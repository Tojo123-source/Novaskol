#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

INPUT_FILE = r'G:\wamp64\www\novaskol-laravel\LIVRE_SOUTENANCE.md'
OUTPUT_DOCX = r'G:\wamp64\www\novaskol-laravel\LIVRE_SOUTENANCE.docx'

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_paragraph_shading(paragraph, color):
    pPr = paragraph._element.get_or_add_pPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    pPr.append(shading)

def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(8)
    p.style.font.color.rgb = RGBColor(150, 150, 150)
    run = p.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar1)
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(150, 150, 150)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._element.append(instrText)
    run3 = p.add_run()
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor(150, 150, 150)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._element.append(fldChar2)

def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.style.font.size = Pt(7)
    run = p.add_run('Livre de Soutenance — Novaskol v1.0.6')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.italic = True

def add_separator_line(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '00C853')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_document():
    doc = Document()

    for style_name in ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        style = doc.styles[style_name]
        style.font.name = 'Calibri'
        rPr = style.element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'Calibri')
        rPr.append(rFonts)

    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    style = doc.styles['Title']
    style.font.size = Pt(24)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_before = Pt(24)
    style.paragraph_format.space_after = Pt(24)

    style = doc.styles['Heading 1']
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x08, 0x0e, 0x18)
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(10)
    style.paragraph_format.page_break_before = True

    style = doc.styles['Heading 2']
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x08, 0x0e, 0x18)
    style.paragraph_format.space_before = Pt(14)
    style.paragraph_format.space_after = Pt(8)

    style = doc.styles['Heading 3']
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x0f, 0x1a, 0x30)
    style.paragraph_format.space_before = Pt(10)
    style.paragraph_format.space_after = Pt(6)

    style = doc.styles['Heading 4']
    style.font.size = Pt(11)
    style.font.bold = True
    style.font.italic = True
    style.paragraph_format.space_before = Pt(8)
    style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    section = doc.sections[0]
    add_header(section)
    add_page_number_footer(section)

    r = doc.add_paragraph()
    r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = r.add_run('LIVRE DE SOUTENANCE')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(8, 14, 24)
    r.paragraph_format.space_before = Pt(80)

    r2 = doc.add_paragraph()
    r2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = r2.add_run('─' * 40)
    run.font.color.rgb = RGBColor(0, 200, 83)
    run.font.size = Pt(10)

    r3 = doc.add_paragraph()
    r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = r3.add_run('Novaskol')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(8, 14, 24)
    r3.paragraph_format.space_before = Pt(10)

    r4 = doc.add_paragraph()
    r4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = r4.add_run("Conception et realisation d'une application de gestion\nscolaire hors-ligne avec Laravel, Electron et SQLite")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)
    r4.paragraph_format.space_after = Pt(50)

    r5 = doc.add_paragraph()
    r5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines_meta = [
        "Presente par : Tojo Nambinina RANDRIAMIFALY",
        "Encadre par : [Nom de l'encadrant]",
        "",
        "Communication en Audiovisuelle et Numerique (CAN)",
        "Specialisation Developpement Web",
        "Annee universitaire 2025-2026",
    ]
    for i, line in enumerate(lines_meta):
        if line == "":
            r5.paragraph_format.space_after = Pt(2)
            r5 = doc.add_paragraph()
            r5.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if i < 2:
            run = r5.add_run(line + '\n')
            run.font.size = Pt(12)
            run.font.bold = True
        else:
            run = r5.add_run(line + '\n')
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()

    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph()
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                for side in ['left', 'right', 'top', 'bottom']:
                    border = OxmlElement(f'w:{side}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:color'), 'E0E0E0')
                    border.set(qn('w:space'), '4')
                    pBdr.append(border)
                pPr.append(pBdr)
                set_paragraph_shading(p, 'F5F5FA')
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(26, 26, 46)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.0
                code_block_lines = []
                i += 1
                continue
            else:
                in_code_block = True
                code_block_lines = []
                i += 1
                continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        if '|' in line and line.strip().startswith('|'):
            table_lines.append(line)
            in_table = True
            i += 1
            continue

        if in_table and (i >= len(lines) or not lines[i].strip().startswith('|')):
            in_table = False
            if len(table_lines) >= 2:
                rows_data = []
                for tbl_line in table_lines:
                    cells = [c.strip() for c in tbl_line.strip().split('|')[1:-1]]
                    rows_data.append(cells)
                header_row = rows_data[0]
                data_rows = []
                for row in rows_data[1:]:
                    if all(re.match(r'^[-:\s]+$', c) for c in row if c):
                        continue
                    data_rows.append(row)
                if data_rows:
                    num_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    tblStyle = table.style
                    tblStyle.font.size = Pt(9)
                    for j, cell_text in enumerate(header_row):
                        if j < num_cols:
                            cell = table.rows[0].cells[j]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(cell_text)
                            run.bold = True
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            set_cell_shading(cell, '080E18')
                    for r_idx, row_data in enumerate(data_rows):
                        for j, cell_text in enumerate(row_data):
                            if j < num_cols:
                                cell = table.rows[r_idx + 1].cells[j]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                run = p.add_run(cell_text)
                                run.font.size = Pt(9)
                                if r_idx % 2 == 0:
                                    set_cell_shading(cell, 'F5F5FA')
                    doc.add_paragraph()
            table_lines = []
            continue

        if in_table:
            i += 1
            continue

        if re.match(r'^-{3,}$', line.strip()):
            if line.strip() == '---' or True:
                p = doc.add_paragraph()
                add_separator_line(p)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if text:
                if level == 1:
                    doc.add_heading(text, level=1)
                elif level == 2:
                    doc.add_heading(text, level=2)
                elif level == 3:
                    doc.add_heading(text, level=3)
                elif level == 4:
                    doc.add_heading(text, level=4)
            i += 1
            continue

        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)

        if not line.strip():
            i += 1
            continue

        ol_match = re.match(r'^(\d+)\.\s+(.+)$', line)
        if ol_match:
            num = ol_match.group(1)
            text = ol_match.group(2)
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            run = p.add_run(f'{num}. ')
            run.bold = True
            run.font.color.rgb = RGBColor(0x08, 0x0e, 0x18)
            p.add_run(text)
            i += 1
            continue

        ul_match = re.match(r'^[\-\*]\s+(.+)$', line)
        if ul_match:
            text = ul_match.group(1)
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            run = p.add_run('  ')
            p.add_run(text)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            i += 1
            continue

        if line.strip().startswith('+') and ('---' in line or '|' in line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(26, 26, 46)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            i += 1
            continue

        if any(c in line for c in ['┌', '┐', '└', '┘', '│', '─', '├', '┤', '┬', '┴', '┼', '▲', '▼', '◄', '►']):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(26, 26, 46)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            i += 1
            continue

        text = line.strip()
        if text:
            p = doc.add_paragraph(text, style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        i += 1

    doc.save(OUTPUT_DOCX)
    print(f'Document sauvegarde: {OUTPUT_DOCX}')
    print(f'Taille: {os.path.getsize(OUTPUT_DOCX) / 1024:.1f} KB')

if __name__ == '__main__':
    create_document()
